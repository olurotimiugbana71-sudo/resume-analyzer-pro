"""
Resume Rewriter - AI-Powered Resume Optimization
Copyright 2026 ApexDynamics Solutions | Built by Rotimi Ugbana
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
import PyPDF2
import io

class ResumeRewriter:
    def __init__(self):
        self.action_verbs = {
            "led": "spearheaded, directed, orchestrated, pioneered, championed",
            "managed": "oversaw, coordinated, supervised, administered, governed",
            "created": "engineered, designed, developed, formulated, architected",
            "increased": "amplified, boosted, accelerated, expanded, maximized",
            "improved": "enhanced, optimized, refined, streamlined, elevated",
            "worked": "collaborated, partnered, executed, contributed, drove",
            "helped": "facilitated, supported, enabled, empowered, assisted",
            "made": "produced, generated, delivered, manufactured, constructed",
            "used": "leveraged, utilized, employed, harnessed, deployed",
            "got": "achieved, attained, secured, obtained, acquired"
        }
        
        self.industry_templates = {
            "Technology": {
                "summary": "Innovative Technology Professional with expertise in {skills}. Proven track record of leveraging cutting-edge technologies to drive business growth and operational excellence. Experienced in {experience_area} with a passion for developing scalable solutions.",
                "bullet": "Leveraged {skills} to {action} {result} through {method}",
                "keywords": ["python", "java", "aws", "cloud", "agile", "scrum", "machine learning", "data analysis", "api", "docker", "kubernetes", "git", "ci/cd", "react", "node.js", "sql", "database"]
            },
            "Finance": {
                "summary": "Results-driven Finance Professional specializing in {skills}. Demonstrated expertise in financial analysis, risk management, and strategic planning. Proven ability to drive revenue growth and optimize financial performance.",
                "bullet": "Applied {skills} to {action} {result} through strategic financial planning",
                "keywords": ["financial analysis", "excel", "forecasting", "budgeting", "risk management", "accounting", "audit", "compliance", "financial modeling", "quickbooks"]
            },
            "Healthcare": {
                "summary": "Compassionate Healthcare Professional with extensive experience in {skills}. Dedicated to delivering exceptional patient care while maintaining compliance with healthcare regulations and industry best practices.",
                "bullet": "Delivered {skills} to {action} {result} while maintaining highest standards of care",
                "keywords": ["patient care", "clinical", "hipaa", "medical records", "healthcare", "patient safety", "treatment planning", "diagnosis", "therapy"]
            },
            "Marketing": {
                "summary": "Creative Marketing Professional proficient in {skills}. Track record of developing data-driven campaigns that increase brand awareness and drive customer engagement.",
                "bullet": "Developed {skills} strategies to {action} {result} through targeted campaigns",
                "keywords": ["seo", "social media", "content marketing", "analytics", "branding", "campaign management", "market research", "email marketing", "ppc"]
            },
            "Sales": {
                "summary": "High-performing Sales Professional expert in {skills}. Consistently exceed targets through strategic relationship building and consultative selling approaches.",
                "bullet": "Utilized {skills} to {action} {result} through strategic client relationships",
                "keywords": ["crm", "salesforce", "negotiation", "lead generation", "pipeline", "account management", "b2b", "b2c", "closing", "prospecting"]
            },
            "General": {
                "summary": "Dynamic Professional with strong background in {skills}. Demonstrated ability to {action} and drive organizational success through effective teamwork and problem-solving.",
                "bullet": "Applied {skills} to {action} {result} through collaborative teamwork",
                "keywords": ["leadership", "communication", "teamwork", "problem solving", "project management", "time management", "critical thinking"]
            }
        }
    
    def extract_text(self, file):
        """Extract text from PDF or DOCX"""
        if file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            return " ".join(page.extract_text() for page in reader.pages)
        elif file.name.endswith('.docx'):
            doc = Document(file)
            return " ".join(p.text for p in doc.paragraphs)
    
    def identify_industry(self, text):
        """Determine industry from resume content"""
        text_lower = text.lower()
        industry_scores = {}
        
        for industry, data in self.industry_templates.items():
            score = sum(1 for kw in data["keywords"] if kw in text_lower)
            industry_scores[industry] = score
        
        best = max(industry_scores, key=industry_scores.get)
        return best if industry_scores[best] > 0 else "General"
    
    def extract_skills(self, text):
        """Extract existing skills from resume"""
        text_lower = text.lower()
        all_keywords = []
        for data in self.industry_templates.values():
            all_keywords.extend(data["keywords"])
        
        found = [kw for kw in all_keywords if kw in text_lower]
        return found
    
    def extract_experience(self, text):
        """Extract experience summary"""
        # Look for years of experience
        years_match = re.search(r'(\d+)[\+]?\s*years', text.lower())
        if years_match:
            return f"{years_match.group(1)}+ years of experience"
        return "proven professional experience"
    
    def rewrite_bullet_points(self, text, industry):
        """Rewrite bullet points with action verbs"""
        lines = text.split('\n')
        rewritten = []
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 20:
                continue
            
            # Skip headers
            if line.isupper() or any(h in line.lower() for h in ['summary', 'education', 'experience', 'skills']):
                continue
            
            line_lower = line.lower()
            
            # Replace weak verbs with strong ones
            for weak, strong in self.action_verbs.items():
                if weak in line_lower:
                    strong_word = strong.split(',')[0].strip()
                    line = line.replace(weak, strong_word, 1)
                    line = line[0].upper() + line[1:]
                    break
            
            # Add industry keywords if missing
            template = self.industry_templates[industry]
            missing_kw = [kw for kw in template["keywords"][:3] if kw not in line_lower]
            
            if missing_kw and len(rewritten) < 5:
                line = f"{line} utilizing {', '.join(missing_kw)}"
            
            rewritten.append(line)
        
        return rewritten[:10]  # Return top 10 bullet points
    
    def generate_improved_summary(self, text, industry, skills):
        """Generate improved professional summary"""
        template = self.industry_templates[industry]
        summary = template["summary"]
        
        # Fill in skills
        if skills:
            summary = summary.replace("{skills}", ", ".join(skills[:5]))
        else:
            summary = summary.replace("{skills}", "core competencies")
        
        # Fill in experience area
        summary = summary.replace("{experience_area}", self.extract_experience(text))
        
        # Fill in action
        summary = summary.replace("{action}", "drive results")
        
        return summary
    
    def rewrite_resume(self, file, industry=None):
        """Main rewrite function"""
        text = self.extract_text(file)
        
        # Auto-detect industry if not provided
        if not industry or industry == "Auto-Detect":
            industry = self.identify_industry(text)
        
        # Extract existing skills
        skills = self.extract_skills(text)
        
        # Generate improved sections
        improved_summary = self.generate_improved_summary(text, industry, skills)
        improved_bullets = self.rewrite_bullet_points(text, industry)
        
        # Get industry keywords
        template = self.industry_templates[industry]
        recommended_keywords = template["keywords"][:10]
        
        return {
            "industry": industry,
            "original_skills": skills,
            "improved_summary": improved_summary,
            "improved_bullets": improved_bullets,
            "recommended_keywords": recommended_keywords,
            "action_verbs_used": list(self.action_verbs.values())[:5]
        }
    
    def generate_docx(self, rewrite_result, original_text):
        """Generate optimized resume as DOCX"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Optimized Resume', 0)
        title.alignment = 1  # Center
        
        # Professional Summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(rewrite_result["improved_summary"])
        
        # Skills Section
        doc.add_heading('Core Competencies', level=1)
        skills_text = ", ".join(rewrite_result["recommended_keywords"][:8])
        doc.add_paragraph(skills_text)
        
        # Experience Bullets
        doc.add_heading('Professional Experience', level=1)
        for bullet in rewrite_result["improved_bullets"]:
            doc.add_paragraph(bullet, style='List Bullet')
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output